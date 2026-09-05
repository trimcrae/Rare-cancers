#!/usr/bin/env python3
"""Build NAT Word uploads from the same assembled source as the preview.

Requires LibreOffice Writer on PATH and python-docx. No scientific prose is
retyped here. Output stamps bind complete input and output bytes; visual review
is a separate required step recorded in submission/verification.json.
"""
from __future__ import annotations

import argparse
import hashlib
from html.parser import HTMLParser
import json
from pathlib import Path
import re
import shutil
import subprocess
import tempfile
import xml.etree.ElementTree as ET
import zipfile

import build_candidate as candidate
import build_submission_docx as word
import build_submission_parts as parts
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Mm

OUT = candidate.BUNDLE / "submission"
NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


class BodyText(HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.in_body = False
        self.text = []

    def handle_starttag(self, tag, attrs):
        if tag == "body":
            self.in_body = True
        if tag in {"p", "h1", "h2", "h3", "li", "td", "th", "br"}:
            self.text.append(" ")

    def handle_endtag(self, tag):
        if tag == "body":
            self.in_body = False
        if tag in {"p", "h1", "h2", "h3", "li", "td", "th"}:
            self.text.append(" ")

    def handle_data(self, data):
        if self.in_body:
            self.text.append(data)


def normal(text):
    return " ".join(text.replace("\u00a0", " ").replace("\u00ad", "").split())


def digest(path):
    return hashlib.sha256(path.read_bytes()).hexdigest()


def convert(html, target, soffice, anonymous=False):
    with tempfile.TemporaryDirectory(prefix="aso-word-") as tmp:
        work = Path(tmp)
        source = work / (target.stem + ".html")
        source.write_text(html, encoding="utf-8", newline="\n")
        result = subprocess.run([
            soffice, "--headless", "--norestore",
            "-env:UserInstallation=" + (work / "profile").as_uri(),
            "--infilter=" + word.HTML_FILTER, "--convert-to", word.DOCX_FILTER,
            "--outdir", str(work), str(source)], capture_output=True,
            text=True, encoding="utf-8", errors="replace", timeout=180)
        made = source.with_suffix(".docx")
        if result.returncode or not made.is_file():
            raise RuntimeError(result.stdout + result.stderr)
        doc = Document(made)
        if "Title" not in doc.styles:
            doc.styles.add_style("Title", WD_STYLE_TYPE.PARAGRAPH)
        for style in doc.styles:
            if style.type == 1:
                style.font.color.rgb = RGBColor(0, 0, 0)
                style.paragraph_format.line_spacing = 2.0
        for section in doc.sections:
            section.page_width, section.page_height = Mm(210), Mm(297)
            section.left_margin = section.right_margin = Mm(25)
            section.top_margin = section.bottom_margin = Mm(25)
        usable_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
        for shape in doc.inline_shapes:
            if shape.width > usable_width:
                shape.height = int(shape.height * usable_width / shape.width)
                shape.width = usable_width
        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.line_spacing = Pt(24)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if paragraph.style.name.startswith("Heading"):
                paragraph.paragraph_format.keep_with_next = True
            if paragraph.text.strip() == "Figure legends" and target.stem.startswith("manuscript"):
                paragraph.paragraph_format.page_break_before = True
            if paragraph._p.findall(".//" + qn("w:drawing")):
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.keep_with_next = True
            if paragraph._p.pPr is not None:
                for border in list(paragraph._p.pPr.findall(qn("w:pBdr"))):
                    paragraph._p.pPr.remove(border)
        if doc.paragraphs:
            doc.paragraphs[0].style = doc.styles["Title"]
            doc.paragraphs[0].paragraph_format.keep_with_next = True
        doc.styles["Title"].font.size = Pt(14)
        for table in doc.tables:
            widths = [22, 34, 11, 20, 13] if len(table.columns) == 5 else [13, 34, 31, 22]
            total_width = usable_width
            table.autofit = False
            for column, percent in zip(table.columns, widths):
                column.width = int(total_width * percent / 100)
            for row_index, row in enumerate(table.rows):
                unique = {id(cell._tc) for cell in row.cells}
                for column_index, cell in enumerate(row.cells):
                    cell.width = total_width if len(unique) == 1 else int(total_width * widths[column_index] / 100)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    props = cell._tc.get_or_add_tcPr()
                    borders = OxmlElement("w:tcBorders")
                    for edge in ("top", "left", "bottom", "right"):
                        border = OxmlElement("w:" + edge)
                        for key, value in {"val": "single", "sz": "4", "color": "D9D9D9"}.items():
                            border.set(qn("w:" + key), value)
                        borders.append(border)
                    props.append(borders)
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), "EAF0F5" if row_index < 2 else "FFFFFF")
                    props.append(shading)
                    margins = OxmlElement("w:tcMar")
                    for edge in ("top", "left", "bottom", "right"):
                        margin = OxmlElement("w:" + edge)
                        margin.set(qn("w:w"), "70")
                        margin.set(qn("w:type"), "dxa")
                        margins.append(margin)
                    props.append(margins)
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.line_spacing = 2.0
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in paragraph.runs:
                            run.font.size = Pt(10.5)
                            run.font.color.rgb = RGBColor(0, 0, 0)
        doc.core_properties.author = "" if anonymous else "Tristan D. McRae"
        doc.core_properties.last_modified_by = ""
        doc.core_properties.comments = ""
        doc.save(target)
    parsed = BodyText()
    parsed.feed(html)
    with zipfile.ZipFile(target) as archive:
        body = ET.fromstring(archive.read("word/document.xml"))
        actual = " ".join("".join(t.text or "" for t in p.findall(".//w:t", NS))
                          for p in body.findall(".//w:p", NS))
        if anonymous:
            identity_text = actual + archive.read("docProps/core.xml").decode("utf-8")
            candidate.renderer._assert_anonymous(identity_text)
    expected = normal("".join(parsed.text))
    actual = normal(actual)
    if actual != expected:
        evidence = {"expected": expected, "actual": actual}
        target.with_suffix(".text-mismatch.json").write_text(
            json.dumps(evidence, ensure_ascii=False, indent=2), encoding="utf-8")
        raise RuntimeError(f"Complete text comparison failed: {target.name}")
    return {"complete_body_text_matches": True, "words_including_references_and_tables": len(actual.split())}


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--soffice", default=shutil.which("soffice"))
    args = parser.parse_args()
    if not args.soffice:
        parser.error("LibreOffice Writer is required; pass --soffice or add it to PATH")
    paper = candidate.paper_config()
    parts._TITLE_PAGE_PARTS = tuple(
        {"label": "Declaration of conflicting interest.", "required": True}
        if part.get("heading") == "Author Disclosure Statement" else part
        for part in parts._TITLE_PAGE_PARTS)
    # The raster is a verified derivative of the same source SVG used in the preview.
    canonical_svg = candidate.MANUSCRIPTS / "figures/aso-multipartner-seam.svg"
    if canonical_svg.read_text(encoding="utf-8") != (candidate.BUNDLE / "aso-multipartner-seam.svg").read_text(encoding="utf-8"):
        raise RuntimeError("Candidate Figure 1 differs from the canonical raster source")
    shutil.copy2(canonical_svg.with_suffix(".png"), candidate.BUNDLE / "aso-multipartner-seam.png")
    body, floats = candidate.renderer.assemble(paper, "manuscript")
    title = re.search(r"^#\s+(.*)$", body, re.M).group(1)
    OUT.mkdir(exist_ok=True)
    outputs = {}
    for anonymous in (False, True):
        text = body
        if anonymous:
            text, applied = candidate.renderer.anonymise(text)
            # This revision has a new contributions sentence; remove initials only.
            text = text.replace("T.D.M.", "The author")
            candidate.renderer._assert_anonymous(text)
        html = candidate.renderer.wrap_manuscript(
            title, candidate.renderer.markdown_to_html(text, floats), paper=paper)
        html = re.sub(r'<h2\b[^>]*>Keywords</h2>\s*(?=<p><strong>Keywords)', '', html)
        html, figures = word._svg_to_png(html, paper)
        path = OUT / ("manuscript-anonymized.docx" if anonymous else "manuscript.docx")
        checks = convert(html, path, args.soffice, anonymous)
        checks.update(word._verify(str(path), figures))
        outputs[path.name] = checks
    for name in ("title-page", "figure-legends"):
        part = parts.PARTS[name]
        blocks = part["blocks"](paper)
        path = OUT / (name + ".docx")
        checks = convert(parts._html(paper, part, blocks), path, args.soffice)
        checks.update(parts._verify(str(path), part, blocks))
        outputs[path.name] = checks
    sources = [candidate.BUNDLE / p for p in (
        "manuscript.md", "fusion-junction-aso-journal-tables.md",
        "fusion-junction-aso-journal-references.md", "fusion-junction-aso-sequences.csv",
        "aso-multipartner-seam.svg", "aso-multipartner-seam.png", "build_candidate.py", "build_uploads.py")]
    sources += [canonical_svg, canonical_svg.with_suffix(".png")]
    sources += [candidate.MANUSCRIPTS / p for p in (
        "build_submission_pdf.py", "build_submission_docx.py", "build_submission_parts.py")]
    stamp = {"schema": "aso-word-upload-build/1",
             "sources": {str(p.relative_to(candidate.REPO)).replace('\\', '/'): digest(p) for p in sources},
             "outputs": {name: {"sha256": digest(OUT / name), **checks} for name, checks in outputs.items()},
             "visual_review": "required separately; build success is not visual verification"}
    (OUT / "word-build-stamp.json").write_text(json.dumps(stamp, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(stamp["outputs"], ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
